from docx import Document
import openpyxl
import os
import requests
import json

# 百度API配置
API_KEY = "pVN5lFZCZGQnxPv88PCr363J"
SECRET_KEY = "D8bggHIOZAXZLMbdMZwEnKNKAsC72Ud6"

def get_access_token():
    """获取百度API的access token"""
    url = "https://aip.baidubce.com/oauth/2.0/token"
    params = {
        "grant_type": "client_credentials",
        "client_id": API_KEY,
        "client_secret": SECRET_KEY
    }
    response = requests.post(url, params=params)
    return str(response.json().get("access_token"))

def generate_text_from_prompt(prompt):
    """使用百度API生成文本"""
    url = f"https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie_speed?access_token={get_access_token()}"
    
    payload = json.dumps({
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ]
    })
    headers = {
        'Content-Type': 'application/json'
    }
    
    try:
        response = requests.post(url, headers=headers, data=payload)
        result = response.json()
        if 'result' in result:
            return result['result']
        else:
            print(f"API返回错误: {result}")
            return None
    except Exception as e:
        print(f"调用API时出错: {e}")
        return None

def get_sheet_safely(workbook, sheet_name):
    """安全地获取sheet，如果不存在则返回None"""
    try:
        if sheet_name in workbook.sheetnames:
            return workbook[sheet_name]
        else:
            print(f"提示：在Excel中未找到sheet：{sheet_name}，将跳过处理")
            return None
    except Exception as e:
        print(f"警告：访问sheet {sheet_name} 时出错：{str(e)}")
        return None

def get_cell_value_safely(row, index):
    """安全地获取单元格的值"""
    try:
        return row[index].value if len(row) > index else None
    except Exception:
        return None

def process_sheet_with_api(sheet):
    """处理使用API生成内容的sheet"""
    field_mappings = {}
    if not sheet:
        return field_mappings

    try:
        for row in sheet.iter_rows(min_row=2):  # 从第二行开始
            field_name = get_cell_value_safely(row, 1)  # B列
            prompt = get_cell_value_safely(row, 4)      # E列 - 生成指令
            
            if field_name and prompt:
                print(f"\n处理字段: {field_name}")
                print(f"使用指令: {prompt}")
                
                # 调用API生成内容
                generated_text = generate_text_from_prompt(prompt)
                if generated_text:
                    marker = f"[{field_name}]"
                    field_mappings[marker] = generated_text
                    print(f"成功生成内容，长度: {len(generated_text)}")
                else:
                    print(f"为字段 {field_name} 生成内容失败")
    except Exception as e:
        print(f"警告：处理sheet时出错：{str(e)}")
    
    return field_mappings

def get_field_mappings(sheet):
    """从sheet中获取直接替换的字段映射"""
    field_mappings = {}
    if not sheet:
        return field_mappings

    try:
        for row in sheet.iter_rows(min_row=2):  # 从第二行开始
            field_name = get_cell_value_safely(row, 1)  # B列
            field_value = get_cell_value_safely(row, 5)  # F列
            
            if field_name and field_value:
                field_name = str(field_name).strip()
                field_value = str(field_value).strip()
                marker = f"[{field_name}]"
                field_mappings[marker] = field_value
                print(f"读取到字段：{marker} -> {field_value}")
    except Exception as e:
        print(f"警告：读取sheet时出错：{str(e)}")
    
    return field_mappings

def replace_text_in_document(doc, replacements):
    """在文档中替换文本"""
    replacements_made = False
    
    # 处理段落
    for paragraph in doc.paragraphs:
        text = paragraph.text
        modified = False
        for marker, value in replacements.items():
            if marker in text:
                text = text.replace(marker, value)
                print(f"替换：{marker} -> {value}")
                modified = True
        if modified:
            paragraph.text = text
            replacements_made = True
    
    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                modified = False
                for marker, value in replacements.items():
                    if marker in text:
                        text = text.replace(marker, value)
                        print(f"在表格中替换：{marker} -> {value}")
                        modified = True
                if modified:
                    cell.text = text
                    replacements_made = True
    
    return replacements_made

def process_document(template_path, excel_path, output_path):
    """处理文档的主函数"""
    print("\n=== 开始处理文档 ===")
    
    try:
        # 检查文件是否存在
        if not os.path.exists(template_path):
            print(f"错误：找不到模板文件 {template_path}")
            return False
        
        if not os.path.exists(excel_path):
            print(f"错误：找不到Excel文件 {excel_path}")
            return False
        
        # 读取文件
        print(f"正在读取Excel文件：{excel_path}")
        wb = openpyxl.load_workbook(excel_path)
        
        print(f"正在读取Word模板：{template_path}")
        doc = Document(template_path)
        
        # 处理所有替换
        all_replacements = {}
        
        # 1. 处理生成指令表
        print("\n处理生成指令表...")
        gen_sheet = get_sheet_safely(wb, "生成指令表")
        if gen_sheet:
            gen_replacements = process_sheet_with_api(gen_sheet)
            all_replacements.update(gen_replacements)
        
        # 2. 处理常规替换sheet
        required_sheets = ["封皮", "招标公告", "投标人须知", "评标办法", "合同条款", "发包人要求"]
        for sheet_name in required_sheets:
            print(f"\n处理sheet：{sheet_name}")
            sheet = get_sheet_safely(wb, sheet_name)
            if sheet:
                sheet_replacements = get_field_mappings(sheet)
                all_replacements.update(sheet_replacements)
        
        # 执行替换
        if all_replacements:
            print("\n开始替换文档中的标记...")
            if replace_text_in_document(doc, all_replacements):
                # 保存文档
                print(f"\n正在保存文档：{output_path}")
                doc.save(output_path)
                print("文档保存成功！")
                return True
            else:
                print("\n未找到需要替换的标记")
                return False
        else:
            print("\n没有找到任何可替换的内容")
            return False
            
    except Exception as e:
        print(f"错误：处理文档时出现异常：{str(e)}")
        return False

def main():
    # 文件路径配置
    template_path = "Template.docx"
    excel_path = "需求对应表v3.xlsx"
    output_path = "Template_Generated.docx"
    
    # 执行处理
    success = process_document(template_path, excel_path, output_path)
    
    if success:
        print("\n=== 文档处理成功完成 ===")
    else:
        print("\n=== 文档处理未完全成功，请检查上述日志 ===")

if __name__ == "__main__":
    main()