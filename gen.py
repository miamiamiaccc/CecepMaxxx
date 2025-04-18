# 根据excel中的指令填写文件
import openpyxl
from docx import Document
import requests
import json
import os

# 百度API配置
API_KEY = "pVN5lFZCZGQnxPv88PCr363J"
SECRET_KEY = "D8bggHIOZAXZLMbdMZwEnKNKAsC72Ud6"

def get_access_token():
    """获取百度API的access token"""
    url = "https://aip.baidubce.com/oauth/2.0/token"
    params = {
        "grant_type": "client_credentials",
        "client_id": API_KEY,
        "client_secret": SECRET_KEY
    }
    response = requests.post(url, params=params)
    return str(response.json().get("access_token"))

def generate_text_from_prompt(prompt):
    """使用百度API生成文本"""
    url = f"https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie_speed?access_token={get_access_token()}"
    
    payload = json.dumps({
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ]
    })
    headers = {
        'Content-Type': 'application/json'
    }
    
    try:
        response = requests.post(url, headers=headers, data=payload)
        result = response.json()
        if 'result' in result:
            return result['result']
        else:
            print(f"API返回错误: {result}")
            return None
    except Exception as e:
        print(f"调用API时出错: {e}")
        return None

def process_document():
    # 读取Excel文件
    try:
        wb = openpyxl.load_workbook('需求对应表v2.xlsx')
        sheet = wb['生成指令表']  # 使用指定的sheet
        print("成功读取Excel文件")
    except Exception as e:
        print(f"读取Excel文件失败: {e}")
        return

    # 读取Word模板
    try:
        doc = Document('Template.docx')
        print("成功读取Word模板")
    except Exception as e:
        print(f"读取Word模板失败: {e}")
        return

    # 创建字段映射
    field_mapping = {}
    for row in sheet.iter_rows(min_row=2):  # 从第二行开始
        field_name = row[1].value  # B列 - 字段名
        prompt = row[4].value      # E列 - 生成指令
        
        if field_name and prompt:
            print(f"\n处理字段: {field_name}")
            print(f"使用指令: {prompt}")
            
            # 调用API生成内容
            generated_text = generate_text_from_prompt(prompt)
            if generated_text:
                field_mapping[f"[{field_name}]"] = generated_text
                print(f"成功生成内容，长度: {len(generated_text)}")
            else:
                print(f"为字段 {field_name} 生成内容失败")

    # 替换文档中的标记
    replacement_count = 0
    
    # 处理段落中的标记
    for paragraph in doc.paragraphs:
        for field, value in field_mapping.items():
            if field in paragraph.text:
                print(f"在段落中替换标记: {field}")
                paragraph.text = paragraph.text.replace(field, value)
                replacement_count += 1

    # 处理表格中的标记
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for field, value in field_mapping.items():
                    if field in cell.text:
                        print(f"在表格中替换标记: {field}")
                        cell.text = cell.text.replace(field, value)
                        replacement_count += 1

    print(f"\n总共完成 {replacement_count} 处替换")

    # 保存新文档
    try:
        doc.save('Template2.docx')
        print("\n成功保存新文档: Template2.docx")
    except Exception as e:
        print(f"保存文档失败: {e}")

if __name__ == "__main__":
    print("=== 开始处理文档 ===")
    process_document()
    print("=== 处理完成 ===")