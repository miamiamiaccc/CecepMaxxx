from docx import Document
import openpyxl
import os

def replace_cover_content():
    print("开始处理文档...")
    
    # 读取Excel文件中的封皮sheet
    try:
        print("正在读取Excel文件...")
        wb = openpyxl.load_workbook('需求对应表v2.xlsx')
        sheet = wb['封皮']
        sheet = wb['招标公告']
        sheet = wb['投标人须知']
        
        # 获取Excel中的替换值（B列是字段名称，F列是字段值）
        replace_values = {}
        for row in sheet.iter_rows(min_row=2):  # 从第二行开始，跳过表头
            field_name = row[1].value    # B列
            field_value = row[5].value   # F列
            if field_name and field_value:  # 确保两个值都不为空
                # 确保值是字符串类型
                field_name = str(field_name).strip()
                field_value = str(field_value).strip()
                replace_values[f"[{field_name}]"] = field_value
                print(f"读取到替换项：[{field_name}] -> {field_value}")
            
    except Exception as e:
        print(f"读取Excel文件失败: {e}")
        return
    
    # 读取Word模板文件
    try:
        if not os.path.exists('招标文件/0采购文件封皮.docx'):
            print("错误：找不到招标文件/0采购文件封皮.docx")
            return
            
        print("正在读取Word模板文件...")
        doc = Document('招标文件/0采购文件封皮.docx')
        
    except Exception as e:
        print(f"读取Word文件失败: {e}")
        return
    
    # 替换文档中的所有占位符
    print("\n开始替换占位符...")
    replacements_made = False
    
    # 替换段落中的占位符
    for paragraph in doc.paragraphs:
        for key, value in replace_values.items():
            if key in paragraph.text:
                print(f"找到占位符：{key}")
                # 完全替换段落中的文本
                paragraph.text = paragraph.text.replace(key, value)
                replacements_made = True
    
    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replace_values.items():
                    if key in cell.text:
                        print(f"在表格中找到占位符：{key}")
                        # 完全替换单元格中的文本
                        cell.text = cell.text.replace(key, value)
                        replacements_made = True
    
    if not replacements_made:
        print("警告：未找到任何可替换的占位符")
    
    # 保存新文档
    try:
        print("\n正在保存新文档...")
        doc.save('生成的招标文件/封皮.doc')
        print("文档已成功保存到：生成的招标文件/封皮.doc")
    except Exception as e:
        print(f"保存文档失败: {e}")
        return

if __name__ == "__main__":
    print("=== 开始执行文档替换程序 ===")
    replace_cover_content()
    print("=== 程序执行完成 ===")