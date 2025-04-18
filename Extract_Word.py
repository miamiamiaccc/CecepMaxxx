# 从Template.docx中提取所有标题和内容，生成多个小文档，每个文档对应一个章节

from docx import Document
import io
import os
from docx.shared import Cm
import json
import requests

# 项目要用正文下加项目格式，除了章节，不支持其他样式

# 设置最大宽度为14厘米
MAX_WIDTH_CM = 14.0

# Define initial version number
version = [0, 0, 0]

# Get file name function
def get_file_name(version, heading_text):
    # 创建 version 的副本，这样不会修改原始列表
    temp_version = version.copy()
    
    # 确保列表不会完全为空
    while len(temp_version) > 1 and temp_version[-1] == 0:
        temp_version.pop()
        
    heading_text = ''.join([c for c in heading_text if c.isalnum() or c in (' ', '-', '_')]).strip()
    return '.'.join(map(str, temp_version)) + '- ' + heading_text

# Update version number function
def update_version(level):
    global version
    while len(version) < 3:
        version.append(0)
    
    if level == '1':
        version[0] += 1
        version[1], version[2] = 0, 0
    elif level == '2':
        version[1] += 1
        version[2] = 0
    elif level == '3':
        version[2] += 1

# Save content as new Word document
def save_content_to_new_doc(content, version, heading_text):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    file_name = get_file_name(version, heading_text) + ".docx"
    new_doc = Document()
    
    for item in content:
        if isinstance(item, str):
            paragraph = new_doc.add_paragraph()
            run = paragraph.add_run(item)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif isinstance(item, tuple) and item[0] == 'table':
            table_data = item[1]
            table = new_doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            for i, row in enumerate(table_data):
                for j, cell_text in enumerate(row):
                    table.cell(i, j).text = cell_text
                    for paragraph in table.cell(i, j).paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            tbl = table._element
            tbl_pr = tbl.tblPr
            tbl_borders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tbl_borders.append(border)
            tbl_pr.append(tbl_borders)
        elif isinstance(item, tuple) and item[0] == 'image':
            image_stream = item[1]
            image_name = item[2]
            width_cm = item[3]
            height_cm = item[4]
            if width_cm > MAX_WIDTH_CM:
                scale_factor = MAX_WIDTH_CM / width_cm
                width_cm = MAX_WIDTH_CM
                height_cm = height_cm * scale_factor
            new_doc.add_paragraph().add_run().add_picture(image_stream, width=Cm(width_cm), height=Cm(height_cm))

        elif isinstance(item, tuple) and item[0] == 'list':
            # Add list paragraph to the new document while retaining its list format
            paragraph = new_doc.add_paragraph(item[1], style='List Paragraph')

    new_doc.save(file_name)

# Function to extract images from runs
def get_image_from_run(run):
    drawing_elements = run._element.xpath('.//a:blip')
    
    if drawing_elements:
        embed = drawing_elements[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        image_part = run.part.related_parts[embed]
        image_data = image_part.blob
        drawing_element = run._element.xpath('.//wp:extent')
        if drawing_element:
            cx = int(drawing_element[0].get('cx'))
            cy = int(drawing_element[0].get('cy'))
            width_cm = cx / 914400 * 2.54
            height_cm = cy / 914400 * 2.54
        else:
            width_cm, height_cm = None, None

        image_stream = io.BytesIO(image_data)
        image_stream.name = os.path.basename(image_part.partname)
        return image_stream, image_stream.name, width_cm, height_cm
    
    return None

# Function to iterate over document elements in order
def iter_block_items(parent):
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.document import Document

    parent_elm = parent.element.body if isinstance(parent, Document) else parent._element
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

# Extract numbered headings and content from Word document
def extract_numbered_headings_and_content(docx_path):
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    from docx.oxml.ns import qn

    doc = Document(docx_path)
    
    content_between_headings = []
    current_level = None
    current_heading_text = None
    
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            para = block
            if para.style.name.startswith('Heading'):
                if current_level is not None:
                    save_content_to_new_doc(content_between_headings, version, current_heading_text)
                    content_between_headings = []
                heading_level = para.style.name.split(' ')[-1]
                current_level = heading_level
                update_version(current_level)
                current_heading_text = para.text
                #不加入标题进入正文
                # content_between_headings.append(para.text)
            else:
                paragraph_text = ""
                for run in para.runs:
                    image_data = get_image_from_run(run)
                    if image_data:
                        image_stream, image_name, width_cm, height_cm = image_data
                        content_between_headings.append(('image', image_stream, image_name, width_cm, height_cm))
                    else:
                        text = run.text.strip()
                        if text:
                            paragraph_text += text

                # Check if paragraph is a list (ordered or unordered) and copy it with its format
                if para.style.name in ['List Paragraph'] or para._element.xpath('.//w:numPr'):
                    # Set the paragraph text and ensure all runs use the '宋体' font
                    list_text = para.text
                    for run in para.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置东亚字体为宋体
                    content_between_headings.append(('list', list_text))
                elif paragraph_text:
                    # Ensure non-list paragraphs also use '宋体'
                    for run in para.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    content_between_headings.append(paragraph_text)
    
        elif isinstance(block, Table):
            table = block
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = ''
                    for cell_paragraph in cell.paragraphs:
                        for cell_run in cell_paragraph.runs:
                            image_stream = get_image_from_run(cell_run)
                            if image_stream:
                                pass
                            else:
                                cell_text += cell_run.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
            content_between_headings.append(('table', table_data))
    
    if content_between_headings:
        save_content_to_new_doc(content_between_headings, version, current_heading_text)

# Example: Extract Word document content including images
docx_file = 'Template.docx'  # Replace with your document path
extract_numbered_headings_and_content(docx_file)

def get_access_token():
    """
    使用 AK，SK 生成鉴权签名（Access Token）
    :return: access_token，或是None(如果错误)
    """
    url = "https://aip.baidubce.com/oauth/2.0/token"
    params = {"grant_type": "client_credentials", "client_id": API_KEY, "client_secret": SECRET_KEY}
    try:
        response = requests.post(url, params=params)
        # 打印响应内容以便调试
        print("Token Response:", response.text)
        
        if response.status_code != 200:
            print(f"Error getting access token. Status code: {response.status_code}")
            return None
            
        result = response.json()
        if "access_token" not in result:
            print("Error: No access_token in response")
            print("Response:", result)
            return None
            
        return str(result["access_token"])
    except Exception as e:
        print(f"Exception in get_access_token: {e}")
        return None

def shorten_text(text):
    # 获取访问令牌
    access_token = get_access_token()
    if not access_token:
        print("Failed to get access token")
        return text[:15]  # 返回原文的前15个字符作为后备方案
        
    url = f"https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie_speed?access_token={access_token}"
    
    payload = json.dumps({
        "messages": [
            {
                "role": "user",
                "content": f"{Prompt_Title}'{text}'"
            }
        ]
    })
    headers = {
        'Content-Type': 'application/json'
    }
    
    try:
        response = requests.request("POST", url, headers=headers, data=payload)
        print("API Response Status:", response.status_code)
        print("API Response Text:", response.text)
        
        if response.status_code != 200:
            print(f"API request failed with status code: {response.status_code}")
            return text[:15]
            
        data = json.loads(response.text.strip())
        
        if 'error_code' in data:
            print(f"API Error: {data.get('error_msg', 'Unknown error')}")
            return text[:15]
            
        if 'result' not in data:
            print("No 'result' in API response")
            print("Full response:", data)
            return text[:15]
            
        result_value = data['result']
        cleaned_title = result_value.replace("。", "")  # 去除中文句号"。"
        
        if key_flag == 1:
            if '★' in text and '★' not in cleaned_title:
                cleaned_title = f"★{cleaned_title}"
            elif '▲' in text and '▲' not in cleaned_title:
                cleaned_title = f"▲{cleaned_title}"
                
        return cleaned_title
        
    except Exception as e:
        print(f"Exception in shorten_text: {e}")
        return text[:15]



