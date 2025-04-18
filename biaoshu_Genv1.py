import openpyxl
import docx
import openai
import os
import requests
import json
import re
from docx.shared import Inches, Cm
from docx import Document
import io

# --------------------------------配置部分--------------------------------------------------
# API配置
API_KEY = "你的百度API Key"
SECRET_KEY = "你的百度密钥"

# 文档格式配置
MAX_WIDTH_CM = 14.0  # 设置Word输出最大宽度为14厘米

# 采购文件生成的提示词配置
Prompt_Title = """你是一位资深的政府采购专家，请将以下采购需求概括为10字以内的简短标题，要求准确、规范、专业："""

Prompt_Answer = """你现在是一位资深的政府采购专家，在审核供应商的响应文件。请按照以下格式回复：
1. 首先以"响应情况：完全满足/部分满足/不满足"开头
2. 然后详细说明响应方案，包括：具体实现方式、技术参数、服务承诺等
3. 语言要专业、严谨，符合政府采购文件的规范要求
4. 确保响应内容完全对应采购需求，不偏离主题

采购需求是："""

Prompt_Content = """你是一位资深的政府采购专家，请针对以下采购需求，按照政府采购文件规范，生成详细的技术规格要求（800字左右）。要求：
1. 符合政府采购文件的规范用语和格式
2. 要求明确、具体、可考核
3. 包含必要的技术参数和标准
4. 注意规避歧视性、限制性条款
5. 适当引用相关标准规范

采购需求："""

# 全局变量配置
key_flag = 1  # 启用★（重要）和▲（可选）标记
level1 = 'heading 1'
level2 = 'heading 2'
last_heading_1 = 1  # 从第一章开始
last_heading_2 = 0
last_heading_3 = 0

# --------------------------------函数定义--------------------------------------------------

def get_access_token():
    """获取百度API访问令牌"""
    url = "https://aip.baidubce.com/oauth/2.0/token"
    params = {"grant_type": "client_credentials", "client_id": API_KEY, "client_secret": SECRET_KEY}
    try:
        response = requests.post(url, params=params)
        if response.status_code == 200:
            return str(response.json().get("access_token"))
        else:
            print(f"获取access_token失败: {response.text}")
            return None
    except Exception as e:
        print(f"获取access_token异常: {e}")
        return None

def generate_procurement_requirements(requirement_text):
    """生成采购需求的详细描述"""
    url = "https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie_speed?access_token=" + get_access_token()
    
    payload = json.dumps({
        "messages": [
            {
                "role": "user",
                "content": f"{Prompt_Content}{requirement_text}"
            }
        ]
    })
    headers = {'Content-Type': 'application/json'}
    
    try:
        response = requests.post(url, headers=headers, data=payload)
        data = json.loads(response.text.strip())
        result = data.get('result', '')
        print(f'生成的采购需求:\n{result}')
        return result
    except Exception as e:
        print(f"生成采购需求异常: {e}")
        return requirement_text

def generate_supplier_response(requirement_text):
    """生成供应商响应内容"""
    url = "https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie_speed?access_token=" + get_access_token()
    
    payload = json.dumps({
        "messages": [
            {
                "role": "user",
                "content": f"{Prompt_Answer}{requirement_text}"
            }
        ]
    })
    headers = {'Content-Type': 'application/json'}
    
    try:
        response = requests.post(url, headers=headers, data=payload)
        data = json.loads(response.text.strip())
        result = data.get('result', '')
        print(f'生成的响应内容:\n{result}')
        return result
    except Exception as e:
        print(f"生成响应内容异常: {e}")
        return "完全满足。" + requirement_text

def create_procurement_document(excel_path, template_path, output_path):
    """创建采购文件"""
    # 加载Excel文件
    wb = openpyxl.load_workbook(excel_path)
    ws = wb.active
    
    # 创建新的Word文档
    doc = Document(template_path)
    
    # 遍历Excel行
    for row in ws.iter_rows(min_row=2):  # 跳过标题行
        if not row[0].value:  # 跳过空行
            continue
            
        # 获取各列的值
        section_type = row[1].value  # B列：章节类型
        requirement = row[2].value   # C列：采购需求
        importance = row[4].value    # E列：重要性标记
        
        if not requirement:
            continue
            
        # 根据章节类型设置标题级别
        if section_type == "标题二级":
            heading_level = 2
            last_heading_2 += 1
            last_heading_3 = 0
        else:
            heading_level = 3
            last_heading_3 += 1
            
        # 生成章节号
        section_number = f"{last_heading_1}.{last_heading_2}" if heading_level == 2 else \
                        f"{last_heading_1}.{last_heading_2}.{last_heading_3}"
                        
        # 生成标题
        title = generate_procurement_requirements(requirement)[:15]  # 限制标题长度
        if importance:
            title = f"{importance} {title}"
            
        # 添加标题
        doc.add_heading(f"{section_number} {title}", level=heading_level)
        
        # 添加采购需求
        doc.add_paragraph("采购需求：").bold = True
        doc.add_paragraph(requirement)
        
        # 生成并添加供应商响应
        response = generate_supplier_response(requirement)
        doc.add_paragraph("供应商响应：").bold = True
        doc.add_paragraph(response)
        
        # 添加分隔行
        doc.add_paragraph()
        
    # 保存文档
    doc.save(output_path)
    print(f"采购文件已生成: {output_path}")

# --------------------------------主程序--------------------------------------------------

def main():
    """主程序入口"""
    print("开始生成采购文件...")
    
    # 配置文件路径
    excel_path = "采购需求表.xlsx"
    template_path = "采购文件模板.docx"
    output_path = "生成的采购文件.docx"
    
    try:
        create_procurement_document(excel_path, template_path, output_path)
        print("采购文件生成完成！")
    except Exception as e:
        print(f"生成采购文件时发生错误: {e}")

if __name__ == "__main__":
    main()